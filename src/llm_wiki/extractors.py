from __future__ import annotations

import importlib.util
from dataclasses import dataclass
from pathlib import Path

from .paths import is_text_file


MAX_EXTRACT_CHARS = 2_000_000


@dataclass(frozen=True)
class ExtractedText:
    text: str
    extractor: str
    warnings: list[str]


def extract_text(path: Path, max_chars: int = MAX_EXTRACT_CHARS) -> ExtractedText:
    suffix = path.suffix.lower()
    if is_text_file(path):
        return extract_plain_text(path, max_chars)
    if suffix == ".pdf":
        return extract_pdf(path, max_chars)
    if suffix == ".docx":
        return extract_docx(path, max_chars)
    return ExtractedText("", "metadata-only", [f"No text extractor registered for {suffix or 'extensionless'} files."])


def extract_plain_text(path: Path, max_chars: int) -> ExtractedText:
    try:
        return ExtractedText(path.read_text(encoding="utf-8", errors="replace")[:max_chars], "text", [])
    except OSError as exc:
        return ExtractedText("", "text-error", [f"Could not read text file: {exc}"])


def extract_pdf(path: Path, max_chars: int) -> ExtractedText:
    if importlib.util.find_spec("pypdf") is None:
        return ExtractedText(
            "",
            "missing-pypdf",
            ["Install optional dependency pypdf to index PDF source text: python -m pip install -e .[extract]"],
        )
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        parts = []
        for index, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception as exc:  # pragma: no cover - depends on malformed PDFs
                parts.append(f"\n[Page {index} extraction failed: {exc}]\n")
                continue
            if text.strip():
                parts.append(f"\n[Page {index}]\n{text}")
            if sum(len(part) for part in parts) >= max_chars:
                break
        return ExtractedText("\n".join(parts)[:max_chars], "pypdf", [])
    except Exception as exc:  # pragma: no cover - depends on parser internals
        return ExtractedText("", "pypdf-error", [f"PDF extraction failed: {exc}"])


def extract_docx(path: Path, max_chars: int) -> ExtractedText:
    if importlib.util.find_spec("docx") is None:
        return ExtractedText(
            "",
            "missing-python-docx",
            ["Install optional dependency python-docx to index DOCX source text: python -m pip install -e .[extract]"],
        )
    try:
        from docx import Document

        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return ExtractedText("\n\n".join(parts)[:max_chars], "python-docx", [])
    except Exception as exc:  # pragma: no cover - depends on parser internals
        return ExtractedText("", "python-docx-error", [f"DOCX extraction failed: {exc}"])
